"""Wiki export — PDF, DOCX, CSV output from vault pages.

Reads all .md files from the vault, parses YAML frontmatter and body,
and writes a single output file per format.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any

from pith import output
from pith.config import PithConfig
from pith.i18n import t


class ExportFormat(str, Enum):
    """Supported export formats."""

    pdf = "pdf"
    docx = "docx"
    csv = "csv"


@dataclass(frozen=True)
class WikiPage:
    """A parsed wiki page with frontmatter and body."""

    title: str
    source: str
    ingested_at: str
    schema: str
    references_legislation: str
    body: str

def _parse_frontmatter(text: str) -> tuple[dict[str, Any], str]:
    """Split a wiki page into frontmatter dict and body text."""
    if not text.startswith("---"):
        return {}, text

    end = text.find("---", 3)
    if end == -1:
        return {}, text

    fm_block = text[3:end].strip()
    body = text[end + 3:].strip()

    meta: dict[str, Any] = {}
    for line in fm_block.splitlines():
        sep = line.find(":")
        if sep == -1:
            continue
        key = line[:sep].strip()
        value = line[sep + 1:].strip()
        meta[key] = value

    return meta, body


def _collect_pages(vault_path: Path) -> list[WikiPage]:
    """Read all .md files from the vault and parse into WikiPage objects."""
    pages: list[WikiPage] = []

    for md_file in sorted(vault_path.glob("**/*.md")):
        text = md_file.read_text(encoding="utf-8")
        meta, body = _parse_frontmatter(text)

        pages.append(WikiPage(
            title=meta.get("title", md_file.stem),
            source=meta.get("source", ""),
            ingested_at=meta.get("ingested_at", ""),
            schema=meta.get("schema", ""),
            references_legislation=meta.get("references_legislation", ""),
            body=body,
        ))

    return pages


def _export_pdf(pages: list[WikiPage], output_path: Path, vault_name: str) -> None:
    """Export wiki pages to a single PDF file using fpdf2."""
    import shutil

    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Load a Unicode font — try system fonts in order of preference
    _FONT_CANDIDATES = [
        "/usr/share/fonts/TTF/DejaVuSans.ttf",           # Arch Linux
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",         # Fedora
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", # Ubuntu/Debian
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",  # macOS
        "C:\\Windows\\Fonts\\arial.ttf",                  # Windows
    ]

    font_path: str | None = None
    for candidate in _FONT_CANDIDATES:
        if Path(candidate).exists():
            font_path = candidate
            break

    if font_path:
        pdf.add_font("Unicode", fname=font_path)
        font_name = "Unicode"
    else:
        # TODO: bundle a Unicode font in the build before launch
        font_name = "Helvetica"

    # Metadata page
    pdf.add_page()
    pdf.set_font(font_name, size=20)
    pdf.cell(0, 12, txt=vault_name, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font_name, size=11)
    export_date = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    pdf.cell(0, 8, txt=export_date, new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, txt=f"{len(pages)} pages", new_x="LMARGIN", new_y="NEXT")

    # One section per wiki page
    for page in pages:
        pdf.add_page()
        pdf.set_font(font_name, size=14)
        pdf.cell(0, 10, txt=page.title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, size=10)
        for paragraph in page.body.split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph:
                pdf.multi_cell(0, 6, txt=paragraph)
                pdf.ln(3)

    pdf.output(str(output_path))

def _export_docx(pages: list[WikiPage], output_path: Path, vault_name: str) -> None:
    """Export wiki pages to a single DOCX file using python-docx."""
    from docx import Document

    doc = Document()

    # Table of contents page
    doc.add_heading(vault_name, level=0)
    for page in pages:
        doc.add_paragraph(page.title, style="List Number")
    doc.add_page_break()

    # One section per wiki page
    for page in pages:
        doc.add_heading(page.title, level=1)
        for paragraph in page.body.split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph:
                doc.add_paragraph(paragraph)

    doc.save(str(output_path))


def _export_csv(pages: list[WikiPage], output_path: Path) -> None:
    """Export wiki pages to CSV using pandas."""
    import pandas as pd

    rows = [
        {
            "title": p.title,
            "source": p.source,
            "ingested_at": p.ingested_at,
            "schema": p.schema,
            "references_legislation": p.references_legislation,
            "body": p.body[:500],
        }
        for p in pages
    ]

    df = pd.DataFrame(rows)
    df.to_csv(output_path, index=False, encoding="utf-8")


def export_wiki(
    config: PithConfig,
    fmt: ExportFormat,
    output_path: Path,
) -> None:
    """Export all wiki pages from the vault to a single file.

    Args:
        config:      Validated PITH configuration.
        fmt:         Target export format (pdf, docx, csv).
        output_path: Destination file path.

    Raises:
        SystemExit: If the vault contains no wiki pages.
    """
    vault_path = config.vault.path or Path.cwd()
    vault_name = vault_path.name

    output.info(t("export.start", fmt=fmt.value))

    pages = _collect_pages(vault_path)
    if not pages:
        output.warning(t("export.no_pages"))
        return

    output_path.parent.mkdir(parents=True, exist_ok=True)

    if fmt is ExportFormat.pdf:
        _export_pdf(pages, output_path, vault_name)
    elif fmt is ExportFormat.docx:
        _export_docx(pages, output_path, vault_name)
    elif fmt is ExportFormat.csv:
        _export_csv(pages, output_path)

    output.success(t("export.complete", path=output_path))
