"""DOCX parser — extraction via python-docx.

Paragraphs are grouped into sections by heading styles.
Tables are extracted separately.
"""

from __future__ import annotations

from pathlib import Path

import docx as python_docx

from pith.i18n import t
from pith.parsers.base import ParsedDocument, ParseError, Section, Table


def _is_heading(style_name: str | None) -> bool:
    """Return True if the paragraph style is a heading."""
    if not style_name:
        return False
    return style_name.startswith("Heading")


def parse_docx(path: Path) -> ParsedDocument:
    """Parse a DOCX file into a ``ParsedDocument``.

    Paragraphs are grouped by heading styles. Body text before the
    first heading is placed in a section with an empty heading.

    Args:
        path: Path to the DOCX file.

    Returns:
        A ``ParsedDocument`` with sections split on headings.

    Raises:
        ParseError: If the file cannot be read.
    """
    try:
        doc = python_docx.Document(str(path))
    except Exception as exc:
        raise ParseError(path, t("parser.docx_read_error", detail=str(exc))) from exc

    sections: list[Section] = []
    current_heading = ""
    body_parts: list[str] = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else None

        if _is_heading(style_name):
            # Flush the previous section.
            if body_parts:
                sections.append(Section(heading=current_heading, body="\n".join(body_parts)))
                body_parts = []
            current_heading = para.text.strip()
        else:
            text = para.text.strip()
            if text:
                body_parts.append(text)

    # Flush final section.
    if body_parts:
        sections.append(Section(heading=current_heading, body="\n".join(body_parts)))

    # Extract tables.
    tables: list[Table] = []
    for tbl in doc.tables:
        rows_data: list[list[str]] = []
        for row in tbl.rows:
            rows_data.append([cell.text.strip() for cell in row.cells])
        if rows_data:
            tables.append(Table(headers=rows_data[0], rows=rows_data[1:]))

    title = path.stem
    if doc.core_properties.title:
        title = doc.core_properties.title

    return ParsedDocument(
        title=title,
        sections=sections,
        tables=tables,
        metadata={
            "source": str(path),
            "format": "docx",
        },
    )
